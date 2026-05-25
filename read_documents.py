#!/usr/bin/env python3
"""Extract text from PDF, DOCX, and DOC files.

Dependencies:
    pip install pypdf python-docx pywin32

    pypdf is required for PDF files. python-docx is optional for DOCX files;
    if it is not installed, this script falls back to basic DOCX text extraction
    with the Python standard library. pywin32 and Microsoft Word are only needed
    for legacy .doc files on Windows.

Examples:
    python read_documents.py manual.pdf report.docx legacy.doc
    python read_documents.py spec.pdf -o extracted.txt
    python read_documents.py reference/stm32G4.pdf --pages 120-125
    python read_documents.py reference/stm32G4.pdf --grep HRTIM --context 8 --max-chars 12000
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree


WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
DOCX_FIXED_TEXT_PARTS = (
    "word/document.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/comments.xml",
)


@dataclass(frozen=True)
class TextBlock:
    title: str
    text: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract text from PDF, DOCX, and DOC files."
    )
    parser.add_argument(
        "inputs",
        nargs="+",
        help="One or more .pdf, .docx, or .doc files to read.",
    )
    parser.add_argument(
        "-o",
        "--output",
        help="Optional output text file. If omitted, prints to stdout.",
    )
    parser.add_argument(
        "--encoding",
        default="utf-8",
        help="Encoding used when writing the output file. Default: utf-8.",
    )
    parser.add_argument(
        "--page-separator",
        default="\n\n--- page break ---\n\n",
        help="Separator inserted between PDF pages.",
    )
    parser.add_argument(
        "--file-separator",
        default="\n\n==========\n\n",
        help="Separator inserted between multiple input files.",
    )
    parser.add_argument(
        "--pages",
        help=(
            "PDF-only 1-based pages or ranges to extract, such as "
            "1,3,10-12. If omitted, all PDF pages are read."
        ),
    )
    parser.add_argument(
        "--grep",
        action="append",
        default=[],
        metavar="PATTERN",
        help=(
            "Regex pattern used to keep only matching lines. Repeat this option "
            "to search for multiple patterns. Matching is case-insensitive by default."
        ),
    )
    parser.add_argument(
        "--case-sensitive",
        action="store_true",
        help="Make --grep matching case-sensitive.",
    )
    parser.add_argument(
        "--context",
        type=int,
        default=0,
        help="Number of lines to include before and after each --grep match.",
    )
    parser.add_argument(
        "--max-chars",
        type=int,
        default=0,
        help="Maximum output characters. Use 0 for no limit.",
    )

    args = parser.parse_args()
    if args.context < 0:
        parser.error("--context must be zero or greater")
    if args.max_chars < 0:
        parser.error("--max-chars must be zero or greater")
    return args


def parse_page_selection(selection: str, page_count: int) -> list[int]:
    pages: list[int] = []
    seen: set[int] = set()

    for raw_part in selection.split(","):
        part = raw_part.strip()
        if not part:
            raise RuntimeError("Empty page item in --pages")

        if "-" in part:
            raw_start, raw_end = part.split("-", 1)
            start = parse_page_number(raw_start.strip(), page_count)
            end = parse_page_number(raw_end.strip(), page_count)
            if start > end:
                raise RuntimeError(f"Invalid page range: {part}")
            candidates = range(start, end + 1)
        else:
            candidates = [parse_page_number(part, page_count)]

        for page in candidates:
            if page not in seen:
                seen.add(page)
                pages.append(page)

    return pages


def parse_page_number(value: str, page_count: int) -> int:
    try:
        page = int(value)
    except ValueError as exc:
        raise RuntimeError(f"Invalid page number: {value}") from exc

    if page < 1 or page > page_count:
        raise RuntimeError(f"Page {page} is outside 1-{page_count}")
    return page


def read_pdf(path: Path, page_separator: str) -> str:
    blocks = read_pdf_pages(path, None)
    return page_separator.join(block.text for block in blocks).strip()


def read_pdf_pages(path: Path, page_selection: str | None) -> list[TextBlock]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency 'pypdf'. Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    page_numbers = (
        parse_page_selection(page_selection, page_count)
        if page_selection
        else list(range(1, page_count + 1))
    )

    blocks: list[TextBlock] = []
    for page_number in page_numbers:
        page = reader.pages[page_number - 1]
        text = (page.extract_text() or "").strip()
        blocks.append(TextBlock(f"page {page_number}/{page_count}", text))

    return blocks


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return read_docx_zip(path)

    document = Document(str(path))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return "\n".join(blocks).strip()


def read_docx_zip(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            part_names = [name for name in DOCX_FIXED_TEXT_PARTS if name in names]
            part_names.extend(
                sorted(
                    name
                    for name in names
                    if name.startswith(("word/header", "word/footer"))
                    and name.endswith(".xml")
                )
            )

            if not part_names:
                return ""

            blocks: list[str] = []
            for part_name in part_names:
                blocks.extend(extract_docx_xml_blocks(archive.read(part_name)))
    except zipfile.BadZipFile as exc:
        raise RuntimeError(f"Invalid DOCX file: {path}") from exc
    except ElementTree.ParseError as exc:
        raise RuntimeError(f"Failed to parse DOCX XML in {path}: {exc}") from exc

    return "\n".join(blocks).strip()


def extract_docx_xml_blocks(xml_content: bytes) -> list[str]:
    root = ElementTree.fromstring(xml_content)
    blocks: list[str] = []

    for paragraph in root.iter(f"{WORD_NAMESPACE}p"):
        text = extract_docx_paragraph_text(paragraph).strip()
        if text:
            blocks.append(text)

    return blocks


def extract_docx_paragraph_text(paragraph: ElementTree.Element) -> str:
    parts: list[str] = []
    for node in paragraph.iter():
        if node.tag == f"{WORD_NAMESPACE}t" and node.text:
            parts.append(node.text)
        elif node.tag == f"{WORD_NAMESPACE}tab":
            parts.append("\t")
        elif node.tag in {f"{WORD_NAMESPACE}br", f"{WORD_NAMESPACE}cr"}:
            parts.append("\n")

    return "".join(parts)


def read_doc(path: Path) -> str:
    if sys.platform != "win32":
        raise RuntimeError(
            "Legacy .doc extraction in this script is only supported on Windows."
        )

    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "Reading .doc requires pywin32. Install it with: pip install pywin32"
        ) from exc

    word = None
    document = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        return document.Content.Text.strip()
    except Exception as exc:
        raise RuntimeError(
            "Failed to read .doc via Microsoft Word automation. "
            "Make sure Microsoft Word is installed and the file is not locked. "
            f"Details: {exc}"
        ) from exc
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def extract_text(path: Path, page_separator: str) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path, page_separator)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".doc":
        return read_doc(path)

    raise RuntimeError(
        f"Unsupported file type: {path.suffix}. Expected .pdf, .docx, or .doc"
    )


def extract_blocks(path: Path, page_selection: str | None) -> list[TextBlock]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf_pages(path, page_selection)
    if page_selection:
        raise RuntimeError("--pages is only supported for PDF files")
    if suffix == ".docx":
        return [TextBlock("document", read_docx(path))]
    if suffix == ".doc":
        return [TextBlock("document", read_doc(path))]

    raise RuntimeError(
        f"Unsupported file type: {path.suffix}. Expected .pdf, .docx, or .doc"
    )


def compile_patterns(patterns: list[str], case_sensitive: bool) -> list[re.Pattern[str]]:
    flags = 0 if case_sensitive else re.IGNORECASE
    try:
        return [re.compile(pattern, flags) for pattern in patterns]
    except re.error as exc:
        raise RuntimeError(f"Invalid --grep pattern: {exc}") from exc


def filter_blocks_by_patterns(
    blocks: list[TextBlock], patterns: list[re.Pattern[str]], context: int
) -> list[TextBlock]:
    filtered: list[TextBlock] = []

    for block in blocks:
        lines = block.text.splitlines()
        match_lines = {
            index
            for index, line in enumerate(lines)
            if any(pattern.search(line) for pattern in patterns)
        }
        if not match_lines:
            continue

        ranges = merge_line_ranges(match_lines, context, len(lines))
        rendered_lines: list[str] = []
        previous_end = 0
        for start, end in ranges:
            if rendered_lines and start > previous_end:
                rendered_lines.append("...")
            for index in range(start, end):
                marker = ">" if index in match_lines else " "
                rendered_lines.append(f"{marker} L{index + 1}: {lines[index]}")
            previous_end = end

        filtered.append(TextBlock(block.title, "\n".join(rendered_lines)))

    return filtered


def merge_line_ranges(
    match_lines: set[int], context: int, line_count: int
) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for index in sorted(match_lines):
        start = max(0, index - context)
        end = min(line_count, index + context + 1)
        if ranges and start <= ranges[-1][1]:
            ranges[-1] = (ranges[-1][0], max(ranges[-1][1], end))
        else:
            ranges.append((start, end))
    return ranges


def render_blocks(blocks: list[TextBlock], page_separator: str) -> str:
    rendered: list[str] = []
    for block in blocks:
        body = block.text if block.text else "[No text extracted]"
        rendered.append(f"--- {block.title} ---\n{body}".rstrip())
    return page_separator.join(rendered).strip()


def limit_output(text: str, max_chars: int) -> str:
    if max_chars == 0 or len(text) <= max_chars:
        return text

    marker = (
        f"\n\n[Output truncated to {max_chars} characters. "
        "Use --pages, --grep, or a smaller --context to narrow the result.]"
    )
    keep = max(0, max_chars - len(marker))
    return text[:keep].rstrip() + marker


def render_result(path: Path, text: str) -> str:
    body = text if text else "[No text extracted]"
    return f"===== {path.name} =====\n{body}".rstrip()


def main() -> int:
    args = parse_args()
    rendered_documents: list[str] = []
    had_error = False

    try:
        patterns = compile_patterns(args.grep, args.case_sensitive)
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        return 1

    for raw_path in args.inputs:
        path = Path(raw_path)
        if not path.is_file():
            print(f"[ERROR] File not found: {path}", file=sys.stderr)
            had_error = True
            continue

        try:
            if args.pages or patterns:
                blocks = extract_blocks(path, args.pages)
                if patterns:
                    blocks = filter_blocks_by_patterns(blocks, patterns, args.context)
                text = render_blocks(blocks, args.page_separator) if blocks else "[No matches]"
            else:
                text = extract_text(path, args.page_separator)
        except Exception as exc:
            print(f"[ERROR] {path}: {exc}", file=sys.stderr)
            had_error = True
            continue

        rendered_documents.append(render_result(path, limit_output(text, args.max_chars)))

    if not rendered_documents:
        return 1

    output_text = args.file_separator.join(rendered_documents) + "\n"

    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(output_text, encoding=args.encoding)
    else:
        sys.stdout.write(output_text)

    return 1 if had_error else 0


if __name__ == "__main__":
    raise SystemExit(main())